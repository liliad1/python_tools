import hashlib
import json
import re
import time
import urllib
from urllib.parse import unquote
import os
import requests
from docx import Document
from docx.opc.oxml import qn
from docx.shared import Inches, Pt

def mkdoc(docName, contents):
    doc = Document()
    # doc.add_picture('//gw.alicdn.com/imgextra/i3/667241583/O1CN01Fwf0dH1NZ3YCWEpUe_!!667241583-0-beehive-scenes.jpg')
    doc.add_heading(contents.get('title'), level=1, )
    doc.add_paragraph(contents.get('summary'))
    doc.add_paragraph(contents.get('richText')[0].get('resource')[0].get('text'))
    # doc.font.name = '宋体'
    # doc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    distance = Inches(0.3)

    sec = doc.sections[0]  # sections对应文档中的“节”

    sec.left_margin = distance  # 以下依次设置左、右、上、下页面边距

    sec.right_margin = distance

    sec.top_margin = distance

    sec.bottom_margin = distance

    sec.page_width = Inches(12)  # 设置页面宽度

    sec.page_height = Inches(20)  # 设置页面高度
    doc.save(docName + '.docx')


BASE_PATH = os.path.dirname(os.path.abspath(__file__))
APPKEY = '12574478'
DATA = '{"source": "darenhome", "type": "h5", "userId": "667241583", "page": 1, "tab": "10004"}'
URL = 'https://h5api.m.taobao.com/h5/mtop.taobao.maserati.darenhome.feed/1.0/'
params = {'jsv': '2.5.6', 'appKey': APPKEY, 't': int(time.time() * 1000),
          'sign': 'FAKE_SIGN_WITH_ANYTHING', 'api': 'mtop.taobao.maserati.darenhome.feed', 'v': '1.0',
          'preventFallback': True,
          'type': 'jsonp', 'dataType': 'jsonp', 'callback': 'mtopjsonp', 'a': '上海', 'b': '悠悠',
          'data': DATA}
headers = {
    'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 9_3_4 like Mac OS X) AppleWebKit/601.1.46 ' + \
                  '(KHTML, like Gecko) Version/9.0 Mobile/13G35 Safari/601.1'
}
try:
    DATA1 = {"contentId": "", "source": "darenhome", "type": "h5",
             "params": "{\"sourcePageName\":\"darenhome\"}", "business_spm": "", "track_params": ""}
    URL1 = 'https://h5api.m.taobao.com/h5/mtop.taobao.beehive.detail.contentservicenewv2/1.0/'
    params1 = {'jsv': '2.5.1', 'appKey': APPKEY, 't': int(time.time() * 1000),
               'sign': 'FAKE_SIGN_WITH_ANYTHING', 'api': 'mtop.taobao.beehive.detail.contentservicenewv2', 'v': '1.0',
               'AntiCreep': True, 'AntiFlood': True, 'timeout': 5000,
               'type': 'jsonp', 'dataType': 'jsonp', 'callback': 'mtopjsonp1', 'a': '上海', 'b': '悠悠',
               'data': ''}
    # get token in first request
    r1 = requests.get(URL, params=params, headers=headers)
    token_with_time = r1.cookies.get('_m_h5_tk')
    token = token_with_time.split('_')[0]
    enc_token = r1.cookies.get('_m_h5_tk_enc')
    # get results in second request
    t2 = str(int(time.time() * 1000))
    c = '&'.join([token, t2, APPKEY, DATA])
    m = hashlib.md5()
    m.update(c.encode('utf-8'))
    params.update({'t': t2, 'sign': m.hexdigest()})
    cookies = {'_m_h5_tk': token_with_time, '_m_h5_tk_enc': enc_token}
    r2 = requests.get(URL, params=params, headers=headers, cookies=cookies)
    items = json.loads(r2.text[11:len(r2.text) - 1]).get('data').get('result').get('data').get('feeds')
    for item in items:
        print(item.get('cover'))
        print(item.get('summary'))
        print(item.get('url'))
        print(item.get('items')[0].get('contentId'))
        DATA1['contentId'] = item.get('items')[0].get('contentId')
        params1.update({'data': json.dumps(DATA1)})
        r3 = requests.get(URL1, params=params1, headers=headers)
        token_with_time1 = r3.cookies.get('_m_h5_tk')
        token1 = token_with_time1.split('_')[0]
        enc_token1 = r3.cookies.get('_m_h5_tk_enc')
        # get results in second request
        t3 = str(int(time.time() * 1000))
        c1 = '&'.join([token1, t3, APPKEY, json.dumps(DATA1)])
        m1 = hashlib.md5()
        m1.update(c1.encode('utf-8'))
        cookies1 = {'_m_h5_tk': token_with_time1, '_m_h5_tk_enc': enc_token1}
        params1.update({'t': t3, 'sign': m1.hexdigest()})
        r4 = requests.get(URL1, params=params1, headers=headers, cookies=cookies1)
        contents = json.loads(r4.text[12:len(r4.text) - 1]).get('data').get('models').get('content')
        mkdoc(contents.get('summary'), contents)
        print(contents)
except Exception as e:
    print(e)


